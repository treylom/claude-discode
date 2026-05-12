from __future__ import annotations

import math
import random
import re
import shutil
import subprocess
from datetime import date, timedelta
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
COMPANY = "NuriFlow Systems"
COMPANY_KO = "누리플로우 시스템즈"
BRAND = "#0EA5E9"

MONTHS = [
    ("2025-05-31", "2025년 5월", "고객 발견과 봄 릴리스 후속 정리", 1.86),
    ("2025-06-30", "2025년 6월", "integration backlog 감축과 Singapore customer success 정착", 1.94),
    ("2025-07-31", "2025년 7월", "FlowDesk Connectors 2.0 출시와 제조 고객 확장", 2.02),
    ("2025-08-31", "2025년 8월", "파트너 enablement와 여름 churn watch", 2.10),
    ("2025-09-30", "2025년 9월", "DocuLens AI beta readiness와 보안 검토", 2.19),
    ("2025-10-31", "2025년 10월", "DocuLens AI 신제품 출시와 conversion motion", 2.28),
    ("2025-11-30", "2025년 11월", "공공 파일럿과 People Operations 정비", 2.35),
    ("2025-12-31", "2025년 12월", "연말 renewals와 security assurance", 2.46),
    ("2026-01-31", "2026년 1월", "InsightBridge launch planning과 FY planning", 2.55),
    ("2026-02-28", "2026년 2월", "InsightBridge 신제품 출시와 finance accounts", 2.63),
    ("2026-03-31", "2026년 3월", "Q1 2026 매출 outperformance와 retention", 2.74),
    ("2026-04-30", "2026년 4월", "NuriFlow Marketplace Preview와 partner template 검증", 2.83),
]

PRODUCTS = ["FlowDesk", "DocuLens AI", "InsightBridge", "LaunchOps", "Marketplace Preview"]
SEARCH_ANCHORS = [
    "회사 미션은 teams should spend less time chasing documents and more time making decisions 라는 문장으로 정리된다.",
    "Q1 2026 매출은 USD 7.92M로 기록되며 InsightBridge launch demand와 expansion pipeline이 함께 기여했다.",
    "신제품 출시 기록은 FlowDesk Connectors 2.0, DocuLens AI, InsightBridge, NuriFlow Marketplace Preview 순서로 추적한다.",
    "HR 정책은 hybrid work, documented decisions, manager accountability, security awareness를 기준으로 운영된다.",
]

CATEGORIES = [
    "brochure",
    "contract",
    "proposal",
    "meeting-quarterly",
    "report-monthly",
    "report-annual",
    "datasheet",
    "marketing",
    "hr-doc",
    "policy",
]

PDF_PLAN = [
    ("brochure", 18),
    ("contract", 22),
    ("proposal", 22),
    ("meeting-quarterly", 18),
    ("report-monthly", 24),
    ("report-annual", 12),
    ("datasheet", 20),
    ("marketing", 20),
    ("hr-doc", 18),
    ("policy", 16),
]

DOCX_FOLDERS = [
    ("10-Contracts", "Master Services Agreement", "contract"),
    ("20-Proposals", "Workflow Transformation Proposal", "proposal"),
    ("30-Meetings", "Meeting Notes", "meeting"),
    ("40-Reports", "Operating Report", "report"),
    ("70-HR", "People Operations Document", "hr"),
    ("90-Policies", "Policy Appendix", "policy"),
]

MD_FOLDERS = ["Inbox/Slack", "Inbox/Kakaotalk", "Inbox/Memos", "Inbox/Threads-Drafts", "_meta/MOCs"]


def words(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9가-힣]+", text))


def slug(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9가-힣]+", "-", text).strip("-")
    return re.sub(r"-+", "-", text)


def typ_escape(text: str) -> str:
    return (
        text.replace("\\", "\\\\")
        .replace("#", "\\#")
        .replace("$", "\\$")
        .replace("[", "\\[")
        .replace("]", "\\]")
    )


def typ_str(text: str) -> str:
    return '"' + text.replace("\\", "\\\\").replace('"', '\\"') + '"'


def month_for(i: int) -> tuple[str, str, str, float]:
    return MONTHS[i % len(MONTHS)]


def paragraph_bank(kind: str, title: str, month_label: str, focus: str, revenue: float, idx: int) -> list[str]:
    product = PRODUCTS[idx % len(PRODUCTS)]
    peer = PRODUCTS[(idx + 2) % len(PRODUCTS)]
    return [
        f"{title} 문서는 {COMPANY_KO}의 2025-05-13부터 2026-05-13까지 운영 기록 중 {month_label} 맥락을 설명한다. 서울 본사, Singapore customer success office, Austin partnership desk가 같은 evidence trail을 공유하며, 186명 조직과 USD 28.4M ARR 규모에 맞는 의사결정 속도를 유지하는 것이 핵심이다. 이번 항목의 초점은 {focus}이며, 고객 신뢰, implementation capacity, operating leverage를 함께 검토한다.",
        f"운영팀은 {product} adoption을 단순 사용량이 아니라 승인 cycle time, exception queue aging, source citation 품질, renewal risk signal로 평가한다. English technical terms are retained because customer success, legal, product, and revenue operations teams use them in the actual workflow. 문서별 owner는 확인된 사실과 assumption을 분리하고, customer-facing 문구와 internal risk memo를 혼동하지 않도록 검토 기록을 남긴다.",
        f"{SEARCH_ANCHORS[idx % len(SEARCH_ANCHORS)]} 이 anchor는 검색 검증을 위해 여러 형식의 파일에 반복되지만, 각 문서는 다른 관점에서 같은 사실을 다룬다. 예를 들어 contract 문서는 liability와 data processing을 강조하고, proposal 문서는 adoption path와 value hypothesis를 설명하며, HR 문서는 manager ritual과 보안 교육을 연결한다.",
        f"{month_label} 월간 revenue marker는 USD {revenue:.2f}M이다. Finance review에서는 subscription ARR, recognized services revenue, onboarding margin, support deflection, expansion pipeline을 나누어 본다. 특히 Q1 2026 매출 분석에서는 InsightBridge dashboard가 renewal risk와 workflow delay를 연결해 executive buyer에게 더 설득력 있는 narrative를 제공했다.",
        f"리스크 관리는 과장된 automation promise를 줄이고 source paragraph, approved workflow rule, reviewer decision을 남기는 방식으로 설계된다. {peer} 관련 항목은 product marketing이 만든 launch copy와 security team이 승인한 architecture note 사이의 차이를 표시한다. 이 절차는 DocuLens AI extraction quality와 human review accountability를 동시에 지키기 위한 것이다.",
        f"실행 계획은 owner, due date, dependency, customer-visible impact로 정리된다. Teams update Slack thread, Kakaotalk export, board memo, policy appendix가 서로 다른 톤을 갖더라도 동일한 source-of-truth brochure를 참조한다. 그래서 사용자는 회사 미션, 신제품 출시, HR 정책, Q1 2026 매출 같은 질문을 여러 파일에서 교차 확인할 수 있다.",
    ]


def body_text(kind: str, title: str, doc_date: str, idx: int, target: int = 520) -> str:
    _, month_label, focus, revenue = month_for(idx)
    sections = [
        ("요약", paragraph_bank(kind, title, month_label, focus, revenue, idx)[:2]),
        ("운영 맥락", paragraph_bank(kind, title, month_label, focus, revenue, idx)[2:4]),
        ("결정 사항", paragraph_bank(kind, title, month_label, focus, revenue, idx)[4:]),
    ]
    text = ""
    for heading, paras in sections:
        text += f"== {heading}\n\n" if kind == "typst" else f"## {heading}\n\n"
        text += "\n\n".join(paras) + "\n\n"
    filler = paragraph_bank(kind, title, month_label, focus, revenue, idx)
    cursor = 0
    while words(text) < target:
        text += filler[cursor % len(filler)] + "\n\n"
        cursor += 1
    if kind == "typst":
        text += "```yaml\nmetric: source-citation-quality\nowner: RevOps Platform\nstatus: reviewed\n```\n"
    return text


def ensure_dirs() -> None:
    for folder in [
        "design/fonts",
        "build",
        "scripts",
        "00-About",
        "10-Contracts",
        "20-Proposals",
        "30-Meetings",
        "40-Reports",
        "70-HR",
        "90-Policies",
        *MD_FOLDERS,
    ]:
        (ROOT / folder).mkdir(parents=True, exist_ok=True)


def write_design() -> None:
    common = f'''#let nf-cyan = rgb("{BRAND}")
#let base-fonts = ("Pretendard", "Noto Sans CJK KR", "Apple SD Gothic Neo", "AppleGothic", "Arial Unicode MS", "Arial")
#let mono-fonts = ("JetBrains Mono", "Sarasa Mono K", "Noto Sans Mono CJK KR", "D2Coding", "Menlo", "Courier")

#let render-doc(meta, body) = {{
  set document(title: meta.title, author: meta.company)
  set page(
    "a4",
    margin: (top: 22mm, bottom: 21mm, left: 19mm, right: 19mm),
    header: align(right)[#text(size: 8pt, fill: gray)[#meta.company · #meta.date]],
    footer: align(center)[#text(size: 8pt, fill: gray)[#context counter(page).display("1")]],
  )
  set text(font: base-fonts, lang: "ko", size: 10.4pt)
  set par(justify: true, leading: 0.65em)
  show heading.where(level: 1): it => block(below: 11pt)[
    #text(size: 20pt, weight: "bold", fill: nf-cyan)[#it.body]
  ]
  show heading.where(level: 2): it => block(above: 8pt, below: 5pt)[
    #text(size: 13pt, weight: "bold")[#it.body]
  ]
  show raw: set text(font: mono-fonts, size: 9pt)
  align(left)[
    = #meta.title
    #text(size: 9pt, fill: gray)[문서 유형: #meta.category · 기준일: #meta.date · 브랜드 컬러: \\#{BRAND[1:]}]
    #line(length: 100%, stroke: nf-cyan + 0.8pt)
    #body
  ]
}}
'''
    (ROOT / "design/common.typ").write_text(common, encoding="utf-8")
    for category in CATEGORIES:
        template = f'''#import "common.typ": render-doc

#let render(meta, body) = {{
  let merged = meta + (category: "{category}")
  render-doc(merged, body)
}}
'''
        (ROOT / f"design/{category}.typ").write_text(template, encoding="utf-8")
    (ROOT / "design/fonts/.gitignore").write_text("*\n!.gitignore\n!README.md\n", encoding="utf-8")
    (ROOT / "design/fonts/README.md").write_text(
        "# Fonts\n\nPlace Pretendard and Noto Sans CJK KR font files here when redistributing the vault. "
        "The Typst templates list those families first, then fall back to Apple SD Gothic Neo/AppleGothic/Arial Unicode MS for local compilation.\n",
        encoding="utf-8",
    )


def write_typ_doc(doc_id: str, category: str, title: str, doc_date: str, idx: int) -> Path:
    folder = ROOT / "build" / doc_id
    folder.mkdir(parents=True, exist_ok=True)
    body = body_text("typst", title, doc_date, idx, target=430 + (idx % 8) * 45)
    body = "\n\n".join(typ_escape(line) if not line.startswith(("=", "```")) else line for line in body.splitlines())
    source = f'''#import "../../design/{category}.typ": render

// ---
// date: {doc_date}
// company: {COMPANY}
// company_ko: {COMPANY_KO}
// category: {category}
// brand_color: {BRAND}
// ---

#let meta = (
  title: {typ_str(title)},
  date: {typ_str(doc_date)},
  company: {typ_str(COMPANY)},
)

#show: doc => render(meta, doc)

{body}
'''
    typ_path = folder / f"{doc_id}.typ"
    typ_path.write_text(source, encoding="utf-8")
    return typ_path


def compile_typ(typ_path: Path) -> None:
    pdf_path = typ_path.with_suffix(".pdf")
    subprocess.run(["typst", "compile", "--root", str(ROOT), str(typ_path), str(pdf_path)], check=True, cwd=ROOT)


def write_docx(path: Path, title: str, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, 0)
    for block in re.split(r"\n{2,}", text):
        clean = re.sub(r"^#+\s*", "", block.strip())
        if clean:
            doc.add_paragraph(clean)
    doc.save(path)


def write_docx_set() -> None:
    for i in range(140):
        folder, suffix, kind = DOCX_FOLDERS[i % len(DOCX_FOLDERS)]
        doc_date, month_label, focus, _ = month_for(i)
        title = f"{month_label} {suffix} {i + 1:03d}"
        content = f"# {title}\n\nDate: {doc_date}\n\n" + body_text(kind, title, doc_date, i, target=360 + (i % 6) * 35)
        write_docx(ROOT / folder / f"{doc_date}-{slug(title)}.docx", title, content)


def write_md_set() -> None:
    labels = ["Slack Export", "Kakaotalk Export", "Decision Memo", "Threads Draft", "Map of Content"]
    for i in range(210):
        folder = MD_FOLDERS[i % len(MD_FOLDERS)]
        doc_date, month_label, focus, _ = month_for(i)
        label = labels[i % len(labels)]
        title = f"{label} {month_label} {i + 1:03d}"
        fm = (
            "---\n"
            f"type: {label.lower().replace(' ', '-')}\n"
            f"date: {doc_date}\n"
            f"company: {COMPANY}\n"
            "tags: [nuriflow, sample-vault, ko-en]\n"
            "---\n\n"
        )
        body = f"# {title}\n\n이 파일은 {focus} 관련 대화와 draft를 보존한다.\n\n" + body_text("markdown", title, doc_date, i, target=320 + (i % 5) * 30)
        (ROOT / folder / f"{doc_date}-{slug(title)}.md").write_text(fm + body, encoding="utf-8")


def write_readme() -> None:
    readme = f"""# NuriFlow Systems Sample Vault

This sample vault models {COMPANY} / {COMPANY_KO}, a Seoul B2B SaaS company with 186 employees, USD 28.4M ARR, HQ Seoul plus Singapore and Austin, and products FlowDesk, DocuLens AI, InsightBridge, LaunchOps, and Marketplace Preview.

## Typst Workflow

Install Typst, then compile all generated sources:

```bash
brew install typst
sample-vault/scripts/build-typst-vault.sh
```

PDFs live in `sample-vault/build/<doc-id>/` beside their `.typ` sources. All generated PDFs in this vault are compiled from Typst sources. `design/common.typ` defines Korean typography with Pretendard and Noto Sans CJK KR at the front of the fallback list, `lang: "ko"`, `par(leading: 0.65em)`, A4 pages, and code block monospace fallbacks.

## Search Checks

The vault intentionally repeats these research anchors across formats: 회사 미션, Q1 2026 매출, 신제품 출시, HR 정책. The dates are distributed from 2025-05-13 through 2026-05-13.
"""
    (ROOT / "README.md").write_text(readme, encoding="utf-8")
    script = """#!/usr/bin/env bash
set -euo pipefail

ROOT="$(cd "$(dirname "${BASH_SOURCE[0]}")/.." && pwd)"
find "$ROOT/build" -name '*.typ' -print0 | sort -z | while IFS= read -r -d '' src; do
  typst compile --root "$ROOT" "$src" "${src%.typ}.pdf"
done
"""
    path = ROOT / "scripts/build-typst-vault.sh"
    path.write_text(script, encoding="utf-8")
    path.chmod(0o755)


def preserve_legacy_pdfs() -> int:
    if not ROOT.exists():
        return 0
    count = 0
    for pdf in ROOT.rglob("*.pdf"):
        if "build" in pdf.relative_to(ROOT).parts:
            continue
        if ".legacy" in pdf.name:
            continue
        legacy = pdf.with_suffix(".legacy.pdf")
        if not legacy.exists():
            pdf.rename(legacy)
            count += 1
    return count


def validate_inventory(legacy_count: int) -> None:
    counts: dict[str, int] = {}
    for p in ROOT.rglob("*"):
        if p.is_file():
            key = p.suffix.lower() or "[none]"
            counts[key] = counts.get(key, 0) + 1
    queries = ["회사 미션", "Q1 2026 매출", "신제품 출시", "HR 정책"]
    hits: dict[str, list[str]] = {}
    text_files = [p for p in ROOT.rglob("*") if p.is_file() and p.suffix.lower() in {".md", ".typ"}]
    for query in queries:
        q_hits = []
        for p in text_files:
            if query in p.read_text(encoding="utf-8", errors="ignore"):
                q_hits.append(str(p.relative_to(ROOT)))
        hits[query] = q_hits[:12]
    lines = [
        "# Generation Report",
        "",
        f"Legacy PDFs renamed: {legacy_count}",
        f"Total files: {sum(counts.values())}",
        "",
        "## Counts",
    ]
    for key in sorted(counts):
        lines.append(f"- {key}: {counts[key]}")
    lines += ["", "## Search Validation"]
    for q in queries:
        lines.append(f"- {q}: {len(hits[q])} sample hits")
        for h in hits[q][:5]:
            lines.append(f"  - {h}")
    (ROOT / "_meta").mkdir(exist_ok=True)
    (ROOT / "_meta/generation-report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    print("\n".join(lines))


def main() -> None:
    random.seed(13)
    legacy_count = preserve_legacy_pdfs()
    ensure_dirs()
    write_design()
    write_readme()

    typ_paths: list[Path] = []
    seq = 0
    for category, count in PDF_PLAN:
        for local in range(count):
            doc_date, month_label, focus, _ = month_for(seq)
            title = f"{month_label} {category.replace('-', ' ').title()} {local + 1:02d}"
            doc_id = f"{doc_date}-{slug(category)}-{local + 1:02d}"
            if seq == 0:
                title = "Company Brochure v1"
                doc_id = "Company-Brochure-v1"
            typ_paths.append(write_typ_doc(doc_id, category, title, doc_date, seq))
            seq += 1

    for typ_path in typ_paths:
        compile_typ(typ_path)

    write_docx_set()
    write_md_set()
    validate_inventory(legacy_count)


if __name__ == "__main__":
    main()
