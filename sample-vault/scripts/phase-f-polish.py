
from __future__ import annotations
import hashlib, re, subprocess, tempfile, shutil
from datetime import date
from pathlib import Path
from docx import Document
ROOT = Path(__file__).resolve().parents[1]
COMPANY='NuriFlow Systems'; MISSION='teams should spend less time chasing documents and more time making decisions'
PRODUCTS=['FlowDesk','DocuLens AI','InsightBridge','LaunchOps','Marketplace Preview']
OFFICES='Seoul HQ, Singapore customer success office, and Austin partnership desk'; EMPLOYEES='186 employees'; ARR='USD 28.4M ARR'
PEOPLE=[('Minseo Park','CEO'),('Jae Kim','VP Product'),('Hana Choi','VP Customer Success'),('Daniel Reyes','VP Engineering'),('Priya Nair','VP Revenue'),('Eunji Lee','Head of People'),('Marcus Hill','Director of Security'),('Sofia Tan','Regional Lead, Singapore'),('Rachel Morgan','Austin Partnerships Lead')]
CUSTOMERS=[('Hanbit Manufacturing','advanced manufacturing'),('BlueHarbor Logistics','regional logistics'),('Mirae Health Network','healthcare operations'),('Northstar Finance','financial services'),('K-Supply Retail','omnichannel retail'),('Aster Public Services','public sector'),('Lumen Education Group','education'),('Vela Energy','renewable energy'),('Daehan Mobility','transportation'),('Atlas Biotech','life sciences')]
def slug(s): return re.sub(r'-+','-',re.sub(r'[^A-Za-z0-9가-힣]+','-',s).strip('-'))
def sint(s,m): return int(hashlib.sha1(s.encode()).hexdigest()[:8],16)%m
def dname(p):
    m=re.search(r'(20\d{2}-\d{2}-\d{2})',p.name); return m.group(1) if m else '2026-05-13'
def month(iso): return date.fromisoformat(iso).strftime('%B %Y')
def meta_for(path, typ, cat, owner, status='approved'):
    return {'Document Type':typ,'Category':cat,'Tags':', '.join(['nuriflow',cat.lower().replace(' ','-'),typ.lower().replace(' ','-')])[:80],'Created Date':dname(path),'Author / Owner':owner,'Status':status}
def add_meta(doc, meta):
    t=doc.add_table(rows=1, cols=2); t.style='Table Grid'; t.cell(0,0).text='Field'; t.cell(0,1).text='Value'
    for k,v in meta.items():
        c=t.add_row().cells; c[0].text=k; c[1].text=v
def write_docx(path,title,typ,cat,owner,sections,status='approved'):
    path.parent.mkdir(parents=True, exist_ok=True); doc=Document(); add_meta(doc, meta_for(path,typ,cat,owner,status)); doc.add_heading(title,0)
    for h,ps in sections:
        doc.add_heading(h,level=1)
        for x in ps: doc.add_paragraph(x)
    doc.save(path)
def rewrite_existing():
    folders={'10-Contracts':('Contract','Legal Operations'),'20-Proposals':('Proposal','Revenue Strategy'),'30-Meetings':('Meeting Notes','Business Operations'),'40-Reports':('Operating Report','Revenue Operations'),'70-HR':('HR Document','People Operations'),'90-Policies':('Policy','Policy Governance')}; out={}
    for folder,(typ,owner) in folders.items():
        n=0
        for path in sorted((ROOT/folder).glob('*.docx')):
            idx=sint(path.name,10000); doc_date=dname(path); product=PRODUCTS[idx%5]; peer=PRODUCTS[(idx+2)%5]; customer,industry=CUSTOMERS[idx%len(CUSTOMERS)]; title=path.stem
            if folder=='10-Contracts':
                ctype=['NDA','Master Services Agreement','Service Schedule','Data Processing Addendum','Partnership Agreement'][idx%5]
                sections=[('Summary',[f'{ctype} for {customer} ({industry}) covering {product} with supporting references to {peer}. The agreement is dated {doc_date} and defines evidence handling, source citation ownership, liability boundaries, data processing rules, and escalation paths for {OFFICES}.',f'Commercial scope assumes {EMPLOYEES}, {ARR}, and the mission that {MISSION}. This version separates customer-facing commitments from internal launch assumptions before signature.']),('Commercial Terms',[f'Primary package: {product} annual subscription, implementation workshop, admin enablement, renewal review, and product-specific acceptance evidence. Liability, confidentiality, and service credits are calibrated to the {ctype}.','Legal Operations owns redlines; Security owns DPA exhibits; Customer Success owns implementation dependencies and acceptance evidence.']),('Risk Controls',['No automated decision is treated as final without a named reviewer, approved workflow rule, and retained source paragraph.',f'Exception requests route through Legal Operations within five business days and are logged against {customer} account governance.'])]
            elif folder=='20-Proposals':
                roi=12+idx%18; sections=[('Summary',[f'Proposal for {customer}, a {industry} organization, to deploy {product} for document intake, review routing, and decision traceability across {OFFICES}.',f'Concrete use case: reduce stalled approval queues by linking every renewal, security review, and customer commitment to the exact source paragraph. Expected ROI is {roi}% within two quarters through faster cycle time and lower manual reconciliation.']),('Customer Scenario',[f'{customer} currently splits document decisions between shared drives, Slack threads, and spreadsheet trackers. {product} becomes the accountable workspace while {peer} supplies cross-team analytics or launch-readiness context.',f'The program supports {COMPANY} mission that {MISSION} while respecting security review and manager accountability.']),('Implementation Path',['Phase 1 maps workflows and access roles; Phase 2 imports priority templates; Phase 3 measures adoption, exception aging, and renewal-risk signal quality.','Revenue Strategy owns value hypothesis, Customer Success owns enablement, and Product owns integration constraints.'])]
            elif folder=='30-Meetings':
                att=[PEOPLE[(idx+i)%len(PEOPLE)] for i in range(3+idx%5)]; att_text='; '.join(f'{a}, {b}' for a,b in att); due1='2026-05-20' if doc_date>='2026-04-01' else '2025-12-15'; due2='2026-06-03' if doc_date>='2026-04-01' else '2026-01-09'
                sections=[('Summary',[f'meeting_date: {doc_date}. Attendees: {att_text}. Topic: {product} operating review for {month(doc_date)} with references to customer evidence, security review, and launch dependencies.',f'Agenda items: adoption trend, unresolved exceptions, {peer} dependency check, customer-facing commitments, and owner readiness for the next review cycle.']),('Decisions',[f'Approved a {product} account-health checkpoint for {CUSTOMERS[(idx+1)%len(CUSTOMERS)][0]} before the next executive update.','Deferred broad automation messaging until source citation quality and reviewer accountability metrics remain above threshold for two consecutive weeks.']),('Action Items',[f'Owner: {att[0][0]} ({att[0][1]}). Due: {due1}. Deliver refreshed risk register with stale exceptions grouped by Sales, CS, Product, and Engineering.',f'Owner: {att[-1][0]} ({att[-1][1]}). Due: {due2}. Publish next meeting packet and attach customer-visible impact notes.','Next meeting: 2026-06-10 10:00 KST unless customer escalation requires an earlier checkpoint.'])]
            elif folder=='40-Reports':
                kpi=3+idx%9; sections=[('Summary',[f'{month(doc_date)} operating report: {product} pipeline quality improved {kpi}% month over month while {peer} reduced unresolved workflow delays by {2+idx%6}%.','Department breakdown: Sales tracked expansion pipeline and win reasons; CS measured onboarding completion and support deflection; Product reviewed feature adoption and citation quality; Engineering monitored integration latency and incident follow-up.']),('KPI Detail',[f'Sales: USD {(1.4+(idx%7)/10):.1f}M qualified expansion. CS: {82+idx%12}% onboarding milestone completion. Product: {64+idx%18}% weekly active admin usage. Engineering: p95 sync latency under {900+idx%300} ms.',f'The report ties {COMPANY} {ARR} baseline to repeatable operating leverage rather than one-off launch activity.']),('Management Actions',['Revenue Operations will compare forecast quality against source evidence before the next pipeline meeting.','Product and Engineering will publish a joint note separating launch copy from verified architecture commitments.'])]
            elif folder=='70-HR':
                aud=['all employees','people managers','new hires','customer-facing teams'][idx%4]; sections=[('Summary',[f'target audience: {aud}. effective_date: {doc_date}. policy owner: People Operations. This HR document explains how {COMPANY} {EMPLOYEES} apply documented decisions while supporting {product} and {peer} programs.','Compliance requirements: complete security awareness, retain source evidence for employment-impacting decisions, and escalate conflicts through manager accountability channels.']),('Specific Clauses',['Clause 1: managers must document the decision basis for hiring, promotion, performance, and access exceptions. Clause 2: employee data may only be used for approved operational purposes.','Clause 3: customer-sensitive examples in training material require anonymization and People Operations approval before reuse.']),('Enforcement and Review',['Enforcement rules include manager remediation, access review, and formal HR escalation for repeated non-compliance.','Review schedule: quarterly owner review with Security and Legal, plus annual refresh before the May operating-planning cycle.'])]
            else:
                domain=['security','data retention','AI governance','partner operations','customer communications'][idx%5]; sections=[('Summary',[f'policy domain: {domain}. scope: {COMPANY}, {OFFICES}, and all product teams supporting {product}. owner: Policy Governance. review cycle: quarterly.','The policy defines responsibilities for drafting, approval, enforcement, exception handling, and evidence retention so teams can make decisions without chasing disconnected documents.']),('Responsibilities and Enforcement',['Business owners maintain procedures; Legal validates obligations; Security approves controls; People Operations trains employees; team leads enforce documented workflows.','Enforcement includes access suspension, corrective training, executive escalation, and customer notice where contract obligations require it.']),('Exceptions and Version History',['Exception process: submit scope, rationale, compensating control, owner, and expiration date to Policy Governance before operational use.',f'Version history: v1.0 baseline on {doc_date}; v1.1 adds {peer} dependency language; next scheduled review on 2026-05-13.'])]
            write_docx(path,title,typ,folder,owner,sections); n+=1
        out[folder]=n
    return out
def tesc(s): return s.replace('\\','\\\\').replace('#','\\#').replace('$','\\$').replace('[','\\[').replace(']','\\]')
def write_pdf(path,title,cat,paras):
    path.parent.mkdir(parents=True, exist_ok=True); template='datasheet' if cat=='50-Products' else 'marketing' if cat=='60-Marketing' else 'report-annual'
    body='\n\n'.join(f'== {tesc(h)}\n\n{tesc(p)}' for h,p in zip(['Overview','Evidence','Next Steps'],paras))
    src_text=f'#import "../../design/{template}.typ": render\n\n#let meta = (title: "{title}", date: "2026-05-13", company: "{COMPANY}")\n#show: doc => render(meta, doc)\n\n{body}\n'
    tmpdir=ROOT/'_meta'/'phase-f-typst'; tmpdir.mkdir(parents=True, exist_ok=True)
    src=tmpdir/(path.stem+'.typ'); src.write_text(src_text,encoding='utf-8')
    subprocess.run(['typst','compile','--root',str(ROOT),str(src),str(path)],check=True,cwd=ROOT)
def add_about():
    docs=[('Company-Overview.docx','Company Overview','Company Overview',[('Profile',[f'{COMPANY} is a Seoul-headquartered B2B SaaS company with {EMPLOYEES}, {ARR}, and offices in Seoul, Singapore, and Austin. Its mission is that {MISSION}.']),('Products',['The portfolio includes FlowDesk, DocuLens AI, InsightBridge, LaunchOps, and Marketplace Preview, each focused on accountable document workflows and decision evidence.']),('Operating Model',['Teams use documented source trails, regional customer success rituals, and product-specific governance to keep customer commitments aligned with implementation capacity.'])]),('Mission-Vision-Values.docx','Mission Vision Values','Company Narrative',[('Mission',[MISSION+'.']),('Vision',['NuriFlow Systems wants every customer-facing team to make auditable decisions from trusted workspaces rather than scattered messages and files.']),('Values',['Evidence before speed; customer commitments are owned; managers document decisions; security is part of every workflow.'])]),('Leadership-Team.docx','Leadership Team','Leadership Profile',[('Executive Team',['Minseo Park, CEO, founded the company after leading enterprise workflow programs. Jae Kim, VP Product, owns FlowDesk and Marketplace Preview. Hana Choi, VP Customer Success, leads Seoul and Singapore adoption. Daniel Reyes, VP Engineering, owns platform reliability. Priya Nair, VP Revenue, manages ARR quality. Eunji Lee, Head of People, owns manager accountability.']),('Leadership Rhythm',['The team reviews product launches, customer escalations, hiring plans, and policy exceptions on a documented quarterly cycle.'])]),('Office-Locations.docx','Office Locations','Location Guide',[('Seoul',['Seoul HQ houses product, engineering, executive leadership, finance, and policy governance.']),('Singapore',['Singapore customer success office supports regional onboarding, regulated-industry enablement, and renewal-risk review.']),('Austin',['Austin partnership desk manages LaunchOps deployments, partner templates, and Marketplace Preview ecosystem outreach.'])]),('Company-History-Timeline.docx','Company History Timeline','Company Timeline',[('Timeline',['2022: FlowDesk prototype shipped. 2023: Seoul HQ expanded. 2024: DocuLens AI pilot launched. 2025: LaunchOps enterprise deployments scaled. 2026: Marketplace Preview partner program opened.']),('Continuity',[f'The mission remained stable: {MISSION}.'])]),('Awards-Recognition.docx','Awards Recognition','Recognition Summary',[('Awards',['2025 Seoul SaaS Workflow Excellence Award; 2025 Asia Customer Evidence Innovation List; 2026 Responsible AI Operations Shortlist; 2026 Partner Workflow Launch Award.']),('Recognition Criteria',['Awards emphasize measurable adoption, evidence traceability, responsible AI review, and customer implementation outcomes.'])]),('Press-Releases.docx','Press Releases','Press Release Collection',[('Quarterly Releases',['Q2 2025 FlowDesk Connectors 2.0 launch; Q3 2025 DocuLens AI beta readiness; Q1 2026 InsightBridge launch; Q2 2026 Marketplace Preview partner beta.']),('Press Contact',['Communications requests route to Revenue Strategy with Legal review for customer claims and launch metrics.'])])]
    out=[]
    for fn,title,typ,sections in docs: write_docx(ROOT/'00-About'/fn,title,typ,'00-About','Executive Operations',sections,'published'); out.append(fn)
    return out
def add_products():
    out=[]
    for product in PRODUCTS:
        ps=product.replace(' ','-'); sections=[('Overview',[f'{product} is part of the {COMPANY} portfolio for accountable document workflows across {OFFICES}. It supports the mission that {MISSION}.']),('Core Capabilities',[f'{product} provides role-based workspaces, source citations, workflow status, exception tracking, and management reporting for regulated customer teams.']),('Metrics',[f'Product KPIs include admin weekly active use, citation quality, approval cycle time, support deflection, and renewal-risk signal coverage against the {ARR} operating baseline.'])]
        for suffix,typ in [('Datasheet','Product Datasheet'),('User-Manual-v1','User Manual'),('Release-Notes-2025-Q4','Release Notes'),('API-Reference','API Reference')]:
            fn=f'{ps}-{suffix}.docx'; write_docx(ROOT/'50-Products'/fn,f'{product} {suffix}',typ,'50-Products','Product Operations',sections,'published'); out.append(fn)
        for suffix in ['Datasheet','User-Manual-v1']:
            fn=f'{ps}-{suffix}.pdf'; write_pdf(ROOT/'50-Products'/fn,f'{product} {suffix}','50-Products',[f'{product} supports document decisions for {COMPANY} customers through governed workflows, source citations, and product-specific ownership.',f'Primary evidence fields include owner, due date, reviewer, customer impact, and source paragraph. LaunchOps and Marketplace Preview are represented in shared portfolio reporting where relevant.','Deployment teams measure adoption, latency, exception aging, and renewal-risk signal quality before expanding the account footprint.']); out.append(fn)
    return out
def add_marketing():
    out=[]
    for fn,title,typ in [('Q1-2026-Marketing-Plan.docx','Q1 2026 Marketing Plan','Marketing Plan'),('Q2-2026-Campaign-Brief.docx','Q2 2026 Campaign Brief','Campaign Brief'),('Social-Media-Calendar-2026.docx','Social Media Calendar 2026','Marketing Calendar'),('Event-Calendar-2026.docx','Event Calendar 2026','Event Calendar')]:
        write_docx(ROOT/'60-Marketing'/fn,title,typ,'60-Marketing','Product Marketing',[('Campaign Focus',[f'Campaigns position FlowDesk, DocuLens AI, InsightBridge, LaunchOps, and Marketplace Preview around the mission that {MISSION}.']),('Audience',['Primary audiences are operations leaders, revenue teams, security reviewers, partner teams, and customer success executives in Korea, Singapore, and the United States.']),('Measurement',['Marketing measures sourced pipeline, product-qualified account movement, launch event attendance, case-study engagement, and sales-cycle acceleration.'])],'published'); out.append(fn)
    for i,product in enumerate(PRODUCTS):
        ps=product.replace(' ','-')
        for customer,industry in CUSTOMERS[i*2:i*2+2]:
            fn=f'{ps}-Case-Study-{slug(customer)}.docx'; write_docx(ROOT/'60-Marketing'/fn,f'{product} Case Study {customer}','Case Study','60-Marketing','Product Marketing',[('Customer Context',[f'{customer} is a {industry} customer using {product} to connect document evidence, workflow ownership, and decision reporting.']),('Outcome',[f'The team reduced approval delay by {18+i*3}% and improved source-citation confidence before renewal review.']),('Quote',[f'{product} helped our team spend less time reconciling disconnected files and more time deciding what to do next.'])],'published'); out.append(fn)
        fn=f'{ps}-White-Paper.pdf'; write_pdf(ROOT/'60-Marketing'/fn,f'{product} White Paper','60-Marketing',[f'{product} addresses document workflow accountability for teams that need auditable decisions across customer, legal, security, and product operations.',f'The white paper links {product} adoption to operating leverage, source citation quality, and launch-readiness evidence for {COMPANY} {ARR} business.','Recommended adoption path: identify top workflow delays, assign owners, preserve source paragraphs, and measure decision velocity by department.']); out.append(fn)
    write_pdf(ROOT/'60-Marketing'/'Brand-Guidelines.pdf','Brand Guidelines','60-Marketing',[f'{COMPANY} brand language centers on accountable document work and the mission that {MISSION}.','Use product names exactly: FlowDesk, DocuLens AI, InsightBridge, LaunchOps, and Marketplace Preview.','Customer claims must cite approved metrics, named owners, and reviewed source evidence.']); out.append('Brand-Guidelines.pdf')
    return out
def add_finance():
    out=[]
    for fn in ['Annual-Financial-Statement-2025.pdf','Investor-Update-2025-Q1.pdf','Investor-Update-2025-Q2.pdf','Investor-Update-2025-Q3.pdf','Investor-Update-2025-Q4.pdf']:
        write_pdf(ROOT/'80-Finance'/fn,fn[:-4].replace('-',' '),'80-Finance',[f'{COMPANY} closed 2025 with {ARR}, {EMPLOYEES}, and offices in Seoul, Singapore, and Austin.','Revenue commentary separates subscription ARR, recognized services revenue, onboarding margin, support deflection, and expansion pipeline.','Product contribution is reviewed across FlowDesk, DocuLens AI, InsightBridge, LaunchOps, and Marketplace Preview.']); out.append(fn)
    for q in range(1,5):
        fn=f'Q{q}-2025-Financial-Report.pdf'; write_pdf(ROOT/'80-Finance'/fn,f'Q{q} 2025 Financial Report','80-Finance',[f'Q{q} 2025 finance review for {COMPANY} connects ARR movement, customer expansion, services margin, and operating expense discipline.','Department views include Sales pipeline, CS retention, Product adoption, and Engineering platform investment.','Management review tracks cash runway, budget variance, and forecast confidence.']); out.append(fn)
    for m in ['2025-10','2025-11','2025-12','2026-01','2026-02']:
        fn=f'Monthly-PnL-Summary-{m}.docx'; write_docx(ROOT/'80-Finance'/fn,f'Monthly PnL Summary {m}','PnL Summary','80-Finance','Finance',[('Revenue',[f'{m} revenue summary ties product ARR to FlowDesk, DocuLens AI, InsightBridge, LaunchOps, and Marketplace Preview account movement.']),('Expenses',['Expense commentary separates Sales, CS, Product, Engineering, G&A, and regional office operating costs.']),('Management View',['Finance flags budget variance, hiring pace, services margin, and forecast risk for the next operating review.'])]); out.append(fn)
    for fn,title in [('Budget-2026.docx','Budget 2026'),('Cap-Table-2026.docx','Cap Table 2026')]:
        write_docx(ROOT/'80-Finance'/fn,title,'Finance Document','80-Finance','Finance',[('Summary',[f'{title} reflects {COMPANY} {ARR}, {EMPLOYEES}, product portfolio, and regional operating footprint.']),('Controls',['Finance owns approvals, Legal reviews equity or contractual obligations, and executive leadership signs final versions.'])]); out.append(fn)
    return out
def first_text(path,limit=500):
    doc=Document(path); return ' '.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])[:limit]
def counts():
    c={p:0 for p in PRODUCTS}
    for path in ROOT.rglob('*'):
        if not path.is_file(): continue
        if path.suffix.lower() in {'.md','.typ'}: txt=path.read_text(encoding='utf-8',errors='ignore')
        elif path.suffix.lower()=='.docx': txt=' '.join(p.text for p in Document(path).paragraphs)
        else: continue
        for p in PRODUCTS: c[p]+=txt.count(p)
    return c
def main():
    logs=[]; rew=rewrite_existing()
    for k,v in rew.items(): logs.append(f'{k}: boilerplate rewrites applied={v}')
    inv={'00-About':add_about(),'50-Products':add_products(),'60-Marketing':add_marketing(),'80-Finance':add_finance()}
    for k,v in inv.items(): logs.append(f'{k}: file count={len(v)}')
    pc=counts(); samples=[]
    for f in ['10-Contracts','20-Proposals','30-Meetings','40-Reports','70-HR','90-Policies']:
        for p in sorted((ROOT/f).glob('*.docx'))[:2]: samples.append((str(p.relative_to(ROOT)),first_text(p)))
    lines=['# Phase F Report','','## Category Completion Log']+[f'- {x}' for x in logs]+['','## New Category Inventory']
    for k,v in inv.items(): lines += [f'### {k} ({len(v)})']+[f'- {x}' for x in v]+['']
    lines += ['## Product Mention Counts']+[f'- {k}: {v}' for k,v in pc.items()]+['','## Rewrite First-500-Character Samples']
    for n,t in samples[:10]: lines += [f'### {n}',t,'']
    (ROOT/'_meta'/'phase-f-report.md').write_text('\n'.join(lines),encoding='utf-8')
    print('\n'.join(logs)); print('Product mention counts:',pc)
if __name__=='__main__': main()
